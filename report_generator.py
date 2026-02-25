from docx import Document
from docx.shared import Pt, Inches, Length
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import datetime
import os


class WordGenerator:
    def __init__(self, filename):
        self.filename = filename
        self.doc = Document()
        self._setup_global_styles()

    def _setup_global_styles(self):
        """
        设置全局样式：
        1. 正文默认为宋体(中文) + Times New Roman(西文)
        2. 字号小四 (12pt)
        3. 行间距 1.5 倍
        """
        style = self.doc.styles['Normal']

        # 设置西文字体
        style.font.name = 'Times New Roman'
        # 强制设置中文字体 (这是核心代码)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置字号
        style.font.size = Pt(12)

        # 设置段落间距
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

    def _add_formatted_text(self, paragraph, text, is_bold=False, font_size=12):
        """
        辅助函数：向段落添加文本，并强制应用中文字体设置。
        防止样式偶尔失效的问题。
        """
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = is_bold
        run.font.name = 'Times New Roman'
        # 再次强制指定中文字体，确保万无一失
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return run

    def generate(self, raw_data, polished_data):
        # --- 1. 标题部分 ---
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 标题字号加大 (三号/16pt)，加粗
        self._add_formatted_text(p, "Weekly Report", is_bold=True, font_size=16)

        # 姓名与日期 (居中或居右，这里按原格式居左，但应用中文样式)
        p_info = self.doc.add_paragraph()
        p_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._add_formatted_text(p_info, f"{raw_data.get('name', '王骄阳')}\n")
        self._add_formatted_text(p_info, datetime.datetime.now().strftime("%Y年%m月%d日"))

        # --- 2. (1) 上周工作计划 ---
        items = raw_data.get('weekly_items', [])
        plan_titles = [item['title'] for item in items]
        self._add_section("(1) 上周工作计划", "上周工作计划：", plan_titles)

        # --- 3. (2) 上周工作小结 (核心: 文字+图片) ---
        title_p = self.doc.add_paragraph()
        self._add_formatted_text(title_p, "(2) 上周工作小结", is_bold=True)

        # 获取AI润色后的内容，如果AI失败则用原始内容
        ai_summaries = polished_data.get('summary_list', [])

        for idx, item in enumerate(items):
            # 1. 写入文字
            text_content = ai_summaries[idx] if idx < len(ai_summaries) else item.get('result', '已完成')
            p_content = self.doc.add_paragraph()
            # 首行缩进两个字符 (约24pt)
            p_content.paragraph_format.first_line_indent = Pt(24)
            self._add_formatted_text(p_content, f"（{idx + 1}）{text_content}")

            # 2. 写入图片
            images = item.get('images', [])
            for img_path in images:
                if os.path.exists(img_path):
                    try:
                        # 图片段落居中
                        p_img = self.doc.add_paragraph()
                        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p_img.add_run()
                        run.add_picture(img_path, width=Inches(5))  # 宽度固定5英寸

                        # 图片标题
                        p_cap = self.doc.add_paragraph()
                        p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        self._add_formatted_text(p_cap, f"图 {idx + 1}-{items.index(item) + 1} 成果展示", font_size=10)
                    except Exception as e:
                        print(f"图片插入失败: {e}")

        # --- 4. 其他板块 ---
        # (3) 本周工作计划
        next_plans = raw_data.get('next_week_plan', [])
        self._add_section("(3) 本周工作计划", "本周工作计划：", next_plans)

        # (4) - (7) 纯文本板块
        long_term = polished_data.get('long_term_goals', raw_data.get('long_term_goals', ''))
        self._add_text_section("(4) 近期内总的工作目标", long_term)

        sig = polished_data.get('significance', raw_data.get('significance', ''))
        self._add_text_section("(5) 上一周和这一周的工作意义", sig)

        lit = polished_data.get('literature', raw_data.get('literature', ''))
        self._add_text_section("(6) 文献阅读情况", lit)

        others = raw_data.get('others', "暂无")
        self._add_text_section("(7) 其他未尽问题", others)

        # 保存文件
        try:
            self.doc.save(self.filename)
            return self.filename
        except PermissionError:
            new_filename = self.filename.replace(".docx", "_new.docx")
            self.doc.save(new_filename)
            return new_filename

    def _add_section(self, title, subtitle, items, is_numbered=True):
        # 标题行
        p_title = self.doc.add_paragraph()
        self._add_formatted_text(p_title, title, is_bold=True)

        # 副标题行
        if subtitle:
            p_sub = self.doc.add_paragraph()
            p_sub.paragraph_format.first_line_indent = Pt(24)
            self._add_formatted_text(p_sub, subtitle)

        # 列表项
        for idx, item in enumerate(items):
            p_item = self.doc.add_paragraph()
            p_item.paragraph_format.first_line_indent = Pt(24)  # 缩进
            prefix = f"（{idx + 1}）" if is_numbered else ""
            self._add_formatted_text(p_item, f"{prefix}{item}")

    def _add_text_section(self, title, content):
        # 标题
        p_title = self.doc.add_paragraph()
        self._add_formatted_text(p_title, title, is_bold=True)

        # 内容 (处理多行)
        if content:
            # 简单处理：如果包含换行符，分割成多段，或者直接作为一段
            # 这里作为一段处理，但保留换行
            p_content = self.doc.add_paragraph()
            p_content.paragraph_format.first_line_indent = Pt(24)
            self._add_formatted_text(p_content, content)